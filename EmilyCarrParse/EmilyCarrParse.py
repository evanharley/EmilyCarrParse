import xmltodict
import pickle
from docx import Document
from docx.shared import Pt
from collections import OrderedDict
from tkinter import simpledialog
from tkinter import messagebox
from tkinter import Tk

class EmilyCarrParser():
    def __init__(self, *args, **kwargs):
        self.data = xmltodict.parse(open('emily-carr.xml', 'rb'))
        self.tk = Tk()

    def _get_accnum_input(self):
        try:
            input = simpledialog.askstring('Input', "Please list accession numbers separated by a ';'", parent = self.tk)
        except:
            messagebox.showerror('ERROR!', "You did not enter any Accession Numbers", parent = self.tk)
            return []
        return input.strip().split('; ')

    def _get_accession_nums(self):
        input = self._get_accnum_input()
        not_in = simpledialog.askstring('Input', "Get accession numbers including previous list? (True/False)", parent = self.tk)
        if not_in == 'True':
            return input
        else:
            accession_nums = self._parse_accession_nums()
            accession_nums_without_input = [num for num in accession_nums if num not in input]
            return accession_nums_without_input

    def _get_lod_input(self):
        input = messagebox.askyesno('Input', 'Do you want to only report a specific Level of Description',
                                       parent = self.tk)
        return input

    def _get_level_of_description(self):
        input = simpledialog.askstring('Input', 'What level of description do you want to report',
                                       parent = self.tk)
        return input
        
    def _parse_level_of_description(self, data, get_input = True):
        if get_input:
            input = self._get_lod_input()
        else:
            input = True
        output = []
        if not input:
            return data
        else:
            if get_input == True:
                input = self._get_level_of_description()
            else:
                input = 'item'
            for item in data:
                if item['@level'] in ['series', 'file'] and input == 'item':
                    try:
                        output.extend(self._parse_level_of_description(item['c'], False))
                    except KeyError:
                        continue
                if item['@level'] == input:
                    output.append(item)
            return output

    def _parse_material_designation(self, data):
        input = simpledialog.askstring('Input', 'What material designation do you want to report?',
                                       parent = self.tk)
        output = []
        if input == '':
            return data
        else:
            for item in data:
                if isinstance(item['controlaccess']['genreform'], list):
                    for thing in item['controlaccess']['genreform']:
                        if thing['#text'] == input:
                            output.append(item)
                else:
                    if item['controlaccess']['genreform']['#text'] == input:
                        output.append(item)
            return output



    def _parse_accession_nums(self):
        data = self.data['ead']['archdesc']['dsc']['c']
        accession_numbers = []
        for row in data:
            has_accession_num = False
            accession_num_loc = row['did']['note']
            if isinstance(accession_num_loc, list):
                for note in accession_num_loc:
                    if isinstance(note['p'], OrderedDict):
                        continue
                    if note['p'].startswith('Accession number'):
                        has_accession_num = True
                        if note['p'][note['p'].find(':')+1:].strip() not in accession_numbers:
                            accession_numbers.append(note['p'][note['p'].find(':')+1:].strip())
                        
            else:
                if isinstance(accession_num_loc['p'], OrderedDict):
                        continue
                if accession_num_loc['p'].startswith('Accession number'):
                    has_accession_num = True
                    if accession_num_loc['p'][accession_num_loc['p'].find(':')+1:].strip() not in accession_numbers:
                        accession_numbers.append(accession_num_loc['p'][accession_num_loc['p'].find(':')+1:].strip())
                        
            if has_accession_num == False:
                accession_numbers.append('None')
        return accession_numbers

    def _gather_unclean_accession_nums(self):
        data = self.data['ead']['archdesc']['dsc']['c']
        identifiers = []
        for row in data:
            accession_num_loc = row['did']['note']
            if isinstance(accession_num_loc, list):
                for note in accession_num_loc:
                    if isinstance(note['p'], OrderedDict):
                        continue
                    if note['p'].startswith('Accession number') and \
                        note['p'].find('E/C/C23') != -1:
                        identifiers.append(row['did']['unitid']['#text'])
                        
            else:
                if isinstance(accession_num_loc['p'], OrderedDict):
                        continue
                if accession_num_loc['p'].startswith('Accession number') and \
                    accession_num_loc['p'].find('E/C/C23') != -1:
                    identifiers.append(row['did']['unitid']['#text'])
        return identifiers

    def _parse(self, data):
        data_dict = {}
        for item in data:
            if item['@level'] =='series':
                data_dict[item['did']['unitid']['#text']] = self.handle_row(item)
                data_dict[item['did']['unitid']['#text']]['children'] = []
                try:
                    for child in item['c']:
                        data_dict[item['did']['unitid']['#text']]['children'].append(self.handle_row(child))
                except KeyError:
                    data_dict[item['did']['unitid']['#text']] = self.handle_row(item)
            if item['@level'] != 'series':    
                data_dict[item['did']['unitid']['#text']] = self.handle_row(item)
        return data_dict

    def handle_row(self, row):
        output_data = {}
        if 'unittitle' in row['did'].keys():
            if isinstance(row['did']['unittitle'], list):
                output_data['title'] = row['did']['unittitle'][0]['#text']
                output_data['oti'] = row['did']['unittitle'][1]['#text']
            else:
                output_data['title'] = row['did']['unittitle']['#text']
        if 'unitdate' in row['did'].keys():
            output_data['date'] = row['did']['unitdate']['#text']
        if 'note' in row['did'].keys():
            notes = row['did']['note']
            for note in notes:
                if isinstance(note, str):
                    continue
                if isinstance(note['p'], OrderedDict):
                    if '#text' not in note['p'].keys():
                        continue
                    if note['p']['#text'].startswith('Exhibited:'):
                        output_data['exhibit_note'] = note['p']['#text']
        if 'unitid' in row['did'].keys():
            output_data['reference_cd'] = row['did']['unitid']['#text']
        if 'physdesc' in row['did'].keys():
            output_data['physdesc'] = row['did']['physdesc']['#text']
        if 'scopecontent' in row.keys():
            if isinstance(row['scopecontent']['p'], OrderedDict):
                output_data['scopecontent'] = row['scopecontent']['p']['#text']
            else:
                output_data['scopecontent'] = row['scopecontent']['p']
        if 'acqinfo' in row.keys():
            output_data['acqinfo'] = row['acqinfo']['p']
        if 'relatedmaterial' in row.keys():
            if isinstance(row['relatedmaterial']['p'], OrderedDict):
                output_data['related'] = row['relatedmaterial']['p']['#text']
            else:
                output_data['related'] = row['relatedmaterial']['p']

        return output_data

    def _split_data_by_accession(self, data):
        stuff = stuff
        return stuff

    def write(self, pkl = False):
        data = self.data['ead']['archdesc']['dsc']['c']
        data = self._parse_level_of_description(data)
        data = self._parse_material_designation(data)
        for report in self._get_accession_nums():
            stuff = []
            no_acc_num = []
            for row in data:
                accession_num_loc = row['did']['note']
                has_accession_num = False
                if isinstance(accession_num_loc, list):

                    for note in accession_num_loc:
                        
                        if isinstance(note['p'], OrderedDict):
                            continue
                        if note['p'].startswith('Accession number'):
                            has_accession_num = True
                            if note['p'].find(report) != -1:
                                stuff.append(row)
                        
                else:
                    if isinstance(accession_num_loc['p'], OrderedDict):
                            continue
                    if accession_num_loc['p'].startswith('Accession number'):
                        has_accession_num = True                    
                        if accession_num_loc['p'].find(report) != -1:
                            stuff.append(row)
                            
                if has_accession_num is False and report == 'None':
                    stuff.append(row)

            if pkl is True:
                return stuff
            data_to_write = self._parse(stuff)
            document = Document()
            for datum in data_to_write.keys():
                if 'children' not in data_to_write[datum].keys():
                    text = self.write_row(data_to_write[datum])
                else:
                    text = self.write_row(data_to_write[datum])
                    for item in data_to_write[datum]['children']:
                        text += '________________________________________________\n'
                        text += '\t'
                        text += self.write_row(item)

                paragraph = document.add_paragraph(text)
                run = paragraph.add_run()
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
            if report.find('/') != -1:
                report = report.replace('/', '_')
            document.save('Emily Carr Fonds-{}.docx'.format(report))
        return 0
                
    def write_row(self, row):
        text = 'Title: ' + row['title']
        if 'date' in row.keys():
            text += ' , Date: ' + row['date'] + '\n '
        else:
            text += '\n '
        if 'oti' in row.keys():
            text += 'Other Title Information: ' + row['oti'] + ' \n '
        if 'reference_cd' in row.keys():
            text += 'Identifier: ' + row['reference_cd'] + ' \n '
        if 'physdesc' in row.keys():
            text += 'Physical Description: ' + row['physdesc'] + ' \n '
        if 'scopecontent' in row.keys():
            text += 'Scope and Content: ' + row['scopecontent'] + ' \n\n '
        if 'acqinfo' in row.keys():
            text += 'Aquisition Information: ' + row['acqinfo'] + ' \n '
        if 'related' in row.keys():
            text += 'Related Items: ' + row['related'] +' \n '
        if 'exhibit_note' in row.keys():
            text += 'Exhibition Notes: ' + row['exhibit_note']

        text += '\n'
        text+='___________________________________________________________________'
        return text



            

if __name__ == '__main__':
    parse = EmilyCarrParser()
    parse.write()
